import csv
import os
import re
from tempfile import NamedTemporaryFile
from pathlib import Path
from typing import Protocol
from xml.etree import ElementTree
from zipfile import ZipFile

from app.config import Settings, get_settings
from app.models import ParsedDocument
from app.services.ocr_service import OCRService, SUPPORTED_IMAGE_SUFFIXES


class OCRTextExtractor(Protocol):
    def extract_text(self, image_path: Path) -> str:
        pass


class DocumentParser:
    text_suffixes = {".txt", ".md", ".csv", ".pdf", ".docx", ".xlsx", ".doc"}
    image_suffixes = SUPPORTED_IMAGE_SUFFIXES
    supported_suffixes = text_suffixes | image_suffixes
    pdf_ocr_max_pages = 50

    def __init__(
        self,
        ocr_service: OCRTextExtractor | None = None,
        settings: Settings | None = None,
    ):
        self._ocr_service = ocr_service
        self.configure(settings or get_settings())

    def configure(self, settings: Settings) -> None:
        self._settings = settings
        self.pdf_max_pages = settings.pdf_max_pages

    def parse(self, file_path: Path, original_filename: str) -> ParsedDocument:
        suffix = file_path.suffix.lower()
        if suffix not in self.supported_suffixes:
            return self._document(
                filename=original_filename,
                status="unsupported",
                full_text="",
                error=f"Unsupported file type: {suffix or 'unknown'}",
            )

        if suffix in self.image_suffixes:
            return self._parse_image(file_path, original_filename)

        try:
            if suffix == ".csv":
                content = self._read_csv(file_path)
            elif suffix == ".pdf":
                content = self._read_pdf(file_path)
            elif suffix == ".docx":
                content = self._read_docx(file_path)
            elif suffix == ".xlsx":
                content = self._read_xlsx(file_path)
            elif suffix == ".doc":
                content = self._read_doc(file_path)
            else:
                content = file_path.read_text(encoding="utf-8", errors="replace")

            return self._document(
                filename=original_filename,
                status="parsed",
                full_text=content,
            )
        except Exception as exc:
            return self._document(
                filename=original_filename,
                status="failed",
                full_text="",
                error=str(exc),
            )

    def _parse_image(self, file_path: Path, original_filename: str) -> ParsedDocument:
        try:
            content = self._get_ocr_service().extract_text(file_path)
            if not content.strip():
                return self._document(
                    filename=original_filename,
                    status="failed",
                    full_text="",
                    error="No text detected in image",
                )

            return self._document(
                filename=original_filename,
                status="ocr_parsed",
                full_text=content,
            )
        except Exception as exc:
            return self._document(
                filename=original_filename,
                status="failed",
                full_text="",
                error=str(exc),
            )

    def _get_ocr_service(self) -> OCRTextExtractor:
        if self._ocr_service is None:
            self._ocr_service = OCRService(
                provider=self._settings.ocr_provider,
                device=self._settings.ocr_device,
                lang=self._settings.ocr_lang,
                enabled=self._settings.ocr_enable,
            )
        return self._ocr_service

    def _read_csv(self, file_path: Path) -> str:
        rows: list[str] = []
        with file_path.open("r", encoding="utf-8", errors="replace", newline="") as csv_file:
            reader = csv.reader(csv_file)
            for row in reader:
                rows.append(", ".join(row))
        return "\n".join(rows)

    def _read_pdf(self, file_path: Path) -> str:
        try:
            import fitz
        except ModuleNotFoundError:
            return self._read_pdf_without_pymupdf(file_path)

        page_texts: list[str] = []
        errors: list[str] = []
        processed_pages = 0
        ocr_pages = 0
        with fitz.open(str(file_path)) as pdf_document:
            total_pages = len(pdf_document)
            page_limit = min(total_pages, self.pdf_max_pages)
            for page_index in range(page_limit):
                page = pdf_document[page_index]
                page_number = page_index + 1
                page_text = page.get_text().strip()
                if len(page_text) < 20:
                    if ocr_pages >= self.pdf_ocr_max_pages:
                        break
                    ocr_pages += 1
                    try:
                        page_text = self._ocr_single_page(page, page_number)
                    except Exception as exc:
                        errors.append(str(exc))
                        page_text = ""
                page_texts.append(f"[Page {page_number}]\n{page_text}")
                processed_pages = page_number

            result = "\n".join(page_texts)
            if processed_pages < total_pages:
                result += f"\n[仅处理前 {processed_pages} 页，共 {total_pages} 页]"

        if any(page_text.split("\n", maxsplit=1)[-1].strip() for page_text in page_texts):
            return result
        if errors:
            raise RuntimeError(f"OCR fallback failed: {errors[0]}")
        return result

    def _extract_pdf_text(self, file_path: Path) -> tuple[str, int]:
        import fitz

        page_texts: list[str] = []
        with fitz.open(str(file_path)) as pdf_document:
            page_count = min(len(pdf_document), self.pdf_max_pages)
            for page_index in range(page_count):
                page = pdf_document[page_index]
                page_number = page_index + 1
                page_texts.append(f"[Page {page_number}]\n{page.get_text()}")
            result = "\n".join(page_texts)
            if page_count < len(pdf_document):
                result += f"\n[仅处理前 {page_count} 页，共 {len(pdf_document)} 页]"
            return result, len(pdf_document)

    def _ocr_pdf_pages(self, file_path: Path) -> str:
        import fitz

        page_texts: list[str] = []
        errors: list[str] = []
        with fitz.open(str(file_path)) as pdf_document:
            page_count = min(len(pdf_document), self.pdf_ocr_max_pages)
            for page_index in range(page_count):
                page = pdf_document[page_index]
                page_number = page_index + 1
                try:
                    page_text = self._ocr_single_page(page, page_number)
                except Exception as exc:
                    errors.append(str(exc))
                    page_text = ""
                page_texts.append(f"[Page {page_number}]\n{page_text}")

            result = "\n".join(page_texts)
            if page_count < len(pdf_document):
                result += f"\n[仅处理前 {page_count} 页，共 {len(pdf_document)} 页]"

        if any(page_text.split("\n", maxsplit=1)[-1].strip() for page_text in page_texts):
            return result
        if errors:
            raise RuntimeError(f"OCR fallback failed: {errors[0]}")
        raise ValueError("No text detected by OCR fallback")

    def _ocr_single_page(self, page, page_number: int) -> str:
        pixmap = page.get_pixmap(dpi=200)
        image_file = NamedTemporaryFile(
            prefix=f"_ocr_page_{page_number}_",
            suffix=".png",
            dir="/tmp",
            delete=False,
        )
        image_path = Path(image_file.name)
        image_file.close()
        try:
            pixmap.save(str(image_path))
            return self._get_ocr_service().extract_text(image_path)
        finally:
            image_path.unlink(missing_ok=True)

    def _read_docx(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ModuleNotFoundError:
            return self._read_docx_without_python_docx(file_path)

        document = Document(file_path)
        sections: list[str] = []

        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        sections.extend(paragraphs)

        for table in document.tables:
            rows = [
                self._format_table_row(cell.text for cell in row.cells)
                for row in table.rows
            ]
            if rows:
                sections.append("\n".join(rows))

        return "\n\n".join(sections)

    def _read_doc(self, file_path: Path) -> str:
        """Parse legacy .doc files via LibreOffice headless conversion."""
        import subprocess
        import tempfile

        libreoffice = os.environ.get("LIBREOFFICE_BIN", "libreoffice")
        with tempfile.TemporaryDirectory() as tmp:
            _ = subprocess.run(
                [
                    libreoffice, "--headless", "--convert-to", "txt:Text",
                    "--outdir", tmp, str(file_path),
                ],
                capture_output=True,
                timeout=120,
            )
            txt = Path(tmp) / f"{file_path.stem}.txt"
            if not txt.exists():
                return ""
            return txt.read_text(encoding="utf-8", errors="replace")

    def _read_xlsx(self, file_path: Path) -> str:
        try:
            from openpyxl import load_workbook
        except ModuleNotFoundError:
            return self._read_xlsx_without_openpyxl(file_path)

        workbook = load_workbook(file_path, read_only=True, data_only=True)
        try:
            sheets: list[str] = []
            for worksheet in workbook.worksheets:
                rows = [
                    self._format_table_row(self._stringify_cell(cell) for cell in row)
                    for row in worksheet.iter_rows(values_only=True)
                    if any(cell is not None for cell in row)
                ]
                sheet_content = f"[Sheet: {worksheet.title}]"
                if rows:
                    sheet_content = f"{sheet_content}\n" + "\n".join(rows)
                sheets.append(sheet_content)
            return "\n\n".join(sheets)
        finally:
            workbook.close()

    def _read_pdf_without_pymupdf(self, file_path: Path) -> str:
        data = file_path.read_bytes()
        if not data.startswith(b"%PDF"):
            raise ValueError("Invalid PDF file")

        text_parts: list[str] = []
        for stream_match in re.finditer(rb"stream\r?\n(.*?)\r?\nendstream", data, re.DOTALL):
            stream = stream_match.group(1)
            for text_match in re.finditer(rb"\((?:\\.|[^\\)])*\)\s*T[Jj]", stream):
                raw_text = text_match.group(0).rsplit(b")", maxsplit=1)[0][1:]
                text_parts.append(self._decode_pdf_literal(raw_text))

        return f"[Page 1]\n{' '.join(text_parts)}" if text_parts else ""

    def _decode_pdf_literal(self, value: bytes) -> str:
        return (
            value.replace(b"\\(", b"(")
            .replace(b"\\)", b")")
            .replace(b"\\\\", b"\\")
            .decode("latin-1", errors="replace")
        )

    def _read_docx_without_python_docx(self, file_path: Path) -> str:
        with ZipFile(file_path) as archive:
            document_root = ElementTree.fromstring(archive.read("word/document.xml"))

        sections: list[str] = []
        body = self._first_xml_child(document_root, "body")
        if body is None:
            return ""

        for child in body:
            if child.tag.endswith("}p"):
                paragraph = self._xml_text(child)
                if paragraph.strip():
                    sections.append(paragraph)
            elif child.tag.endswith("}tbl"):
                rows = []
                for row in self._xml_children(child, "tr"):
                    rows.append(self._format_table_row(self._xml_text(cell) for cell in self._xml_children(row, "tc")))
                if rows:
                    sections.append("\n".join(rows))

        return "\n\n".join(sections)

    def _read_xlsx_without_openpyxl(self, file_path: Path) -> str:
        with ZipFile(file_path) as archive:
            workbook_root = ElementTree.fromstring(archive.read("xl/workbook.xml"))
            rels_root = ElementTree.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
            rels = {relation.attrib["Id"]: relation.attrib["Target"] for relation in rels_root}
            shared_strings = self._read_xlsx_shared_strings(archive)

            sheets: list[str] = []
            for sheet in workbook_root.findall(".//{*}sheet"):
                title = sheet.attrib["name"]
                relationship_id = sheet.attrib["{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"]
                target = rels[relationship_id]
                sheet_path = f"xl/{target}" if not target.startswith("/") else target[1:]
                sheet_root = ElementTree.fromstring(archive.read(sheet_path))
                rows = self._xlsx_rows(sheet_root, shared_strings)
                sheet_content = f"[Sheet: {title}]"
                if rows:
                    sheet_content = f"{sheet_content}\n" + "\n".join(rows)
                sheets.append(sheet_content)

        return "\n\n".join(sheets)

    def _read_xlsx_shared_strings(self, archive: ZipFile) -> list[str]:
        try:
            root = ElementTree.fromstring(archive.read("xl/sharedStrings.xml"))
        except KeyError:
            return []
        return [self._xml_text(item) for item in root.findall("{*}si")]

    def _xlsx_rows(self, sheet_root: ElementTree.Element, shared_strings: list[str]) -> list[str]:
        rows: list[str] = []
        for row in sheet_root.findall(".//{*}row"):
            values = [self._xlsx_cell_value(cell, shared_strings) for cell in row.findall("{*}c")]
            if any(value for value in values):
                rows.append(self._format_table_row(values))
        return rows

    def _xlsx_cell_value(self, cell: ElementTree.Element, shared_strings: list[str]) -> str:
        cell_type = cell.attrib.get("t")
        if cell_type == "inlineStr":
            return self._xml_text(cell)

        value = self._first_xml_child(cell, "v")
        if value is None or value.text is None:
            return ""
        if cell_type == "s":
            return shared_strings[int(value.text)]
        return value.text

    def _first_xml_child(self, element: ElementTree.Element, name: str) -> ElementTree.Element | None:
        for child in element.iter():
            if child.tag.endswith(f"}}{name}"):
                return child
        return None

    def _xml_children(self, element: ElementTree.Element, name: str) -> list[ElementTree.Element]:
        return [child for child in element if child.tag.endswith(f"}}{name}")]

    def _xml_text(self, element: ElementTree.Element) -> str:
        return "".join(text_node.text or "" for text_node in element.iter() if text_node.tag.endswith("}t"))

    def _format_table_row(self, values) -> str:
        return "| " + " | ".join(self._stringify_cell(value) for value in values) + " |"

    def _stringify_cell(self, value) -> str:
        if value is None:
            return ""
        return str(value).replace("\n", " ").replace("|", "\\|").strip()

    def _preview(self, content: str, max_length: int = 1000) -> str:
        normalized = content.strip()
        if len(normalized) <= max_length:
            return normalized
        return f"{normalized[:max_length]}..."

    def _document(
        self,
        *,
        filename: str,
        status: str,
        full_text: str,
        error: str | None = None,
    ) -> ParsedDocument:
        normalized = full_text.strip()
        return ParsedDocument(
            filename=filename,
            status=status,
            full_text=normalized,
            preview=self._preview(normalized),
            error=error,
            original_chars=len(normalized),
            used_chars=len(normalized),
        )
