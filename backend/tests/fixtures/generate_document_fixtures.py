from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import html

FIXTURES_DIR = Path(__file__).parent


def generate_pdf(path: Path = FIXTURES_DIR / "sample.pdf") -> None:
    try:
        import fitz
    except ModuleNotFoundError:
        _generate_pdf_stdlib(path)
        return

    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Audit PDF sample procurement control")
    page.insert_text((72, 92), "Approval evidence includes budget owner review and sign-off.")
    document.save(path)
    document.close()


def _generate_pdf_stdlib(path: Path) -> None:
    stream = (
        b"BT /F1 12 Tf 72 720 Td (Audit PDF sample procurement control) Tj "
        b"0 -20 Td (Approval evidence includes budget owner review and sign-off.) Tj ET"
    )
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(body)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(content)


def generate_scanned_pdf(path: Path = FIXTURES_DIR / "scanned_sample.pdf") -> None:
    try:
        import fitz
    except ModuleNotFoundError:
        _generate_blank_pdf_stdlib(path)
        return

    document = fitz.open()
    page = document.new_page(width=300, height=200)
    pixmap = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 120, 60), False)
    pixmap.clear_with(255)
    page.insert_image(fitz.Rect(72, 72, 192, 132), pixmap=pixmap)
    document.save(path)
    document.close()


def _generate_blank_pdf_stdlib(path: Path) -> None:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 200] /Resources << >> >>",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(body)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(content)


def generate_docx(path: Path = FIXTURES_DIR / "sample.docx") -> None:
    try:
        from docx import Document
    except ModuleNotFoundError:
        _generate_docx_stdlib(path)
        return

    document = Document()
    document.add_paragraph("Audit DOCX sample procurement policy")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Approval"
    table.cell(1, 1).text = "Complete"
    document.save(path)


def _generate_docx_stdlib(path: Path) -> None:
    document_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>Audit DOCX sample procurement policy</w:t></w:r></w:p>
    <w:tbl>
      <w:tr><w:tc><w:p><w:r><w:t>Control</w:t></w:r></w:p></w:tc><w:tc><w:p><w:r><w:t>Status</w:t></w:r></w:p></w:tc></w:tr>
      <w:tr><w:tc><w:p><w:r><w:t>Approval</w:t></w:r></w:p></w:tc><w:tc><w:p><w:r><w:t>Complete</w:t></w:r></w:p></w:tc></w:tr>
    </w:tbl>
    <w:sectPr/>
  </w:body>
</w:document>
'''
    styles_xml = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/></w:style>
</w:styles>
'''
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
  <Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>
</Types>
''')
        archive.writestr("_rels/.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>
</Relationships>
''')
        archive.writestr("word/_rels/document.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>
''')
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/styles.xml", styles_xml)
        archive.writestr("docProps/core.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:title>sample</dc:title></cp:coreProperties>
''')
        archive.writestr("docProps/app.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>tests</Application></Properties>
''')


def generate_xlsx(path: Path = FIXTURES_DIR / "sample.xlsx") -> None:
    try:
        from openpyxl import Workbook
    except ModuleNotFoundError:
        _generate_xlsx_stdlib(path)
        return

    workbook = Workbook()
    first = workbook.active
    first.title = "Procurement"
    first.append(["Project", "Amount", "Status"])
    first.append(["Office supplies", 1200, "Approved"])
    second = workbook.create_sheet("Risks")
    second.append(["Risk", "Level"])
    second.append(["Missing approval", "High"])
    workbook.save(path)


def _sheet_xml(rows: list[list[object]]) -> str:
    row_xml = []
    for row_number, row in enumerate(rows, start=1):
        cells = []
        for column_number, value in enumerate(row, start=1):
            cell_ref = f"{chr(64 + column_number)}{row_number}"
            if isinstance(value, int | float):
                cells.append(f'<c r="{cell_ref}"><v>{value}</v></c>')
            else:
                cells.append(
                    f'<c r="{cell_ref}" t="inlineStr"><is><t>{html.escape(str(value))}</t></is></c>'
                )
        row_xml.append(f'<row r="{row_number}">{"".join(cells)}</row>')
    return '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>{rows}</sheetData>
</worksheet>
'''.format(rows="".join(row_xml))


def _generate_xlsx_stdlib(path: Path) -> None:
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
  <Override PartName="/xl/worksheets/sheet2.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
  <Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>
</Types>
''')
        archive.writestr("_rels/.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>
''')
        archive.writestr("xl/workbook.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets>
    <sheet name="Procurement" sheetId="1" r:id="rId1"/>
    <sheet name="Risks" sheetId="2" r:id="rId2"/>
  </sheets>
</workbook>
''')
        archive.writestr("xl/_rels/workbook.xml.rels", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
  <Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet2.xml"/>
  <Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>
</Relationships>
''')
        archive.writestr("xl/worksheets/sheet1.xml", _sheet_xml([
            ["Project", "Amount", "Status"],
            ["Office supplies", 1200, "Approved"],
        ]))
        archive.writestr("xl/worksheets/sheet2.xml", _sheet_xml([
            ["Risk", "Level"],
            ["Missing approval", "High"],
        ]))
        archive.writestr("xl/styles.xml", '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <fonts count="1"><font><sz val="11"/><name val="Calibri"/></font></fonts>
  <fills count="2"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill></fills>
  <borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>
  <cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>
  <cellXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/></cellXfs>
  <cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>
</styleSheet>
''')


def generate_all() -> None:
    FIXTURES_DIR.mkdir(parents=True, exist_ok=True)
    generate_pdf()
    generate_scanned_pdf()
    generate_docx()
    generate_xlsx()
    (FIXTURES_DIR / "unsupported.doc").write_bytes(b"")
    (FIXTURES_DIR / "unsupported.xls").write_bytes(b"")


if __name__ == "__main__":
    generate_all()
