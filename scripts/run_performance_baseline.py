#!/usr/bin/env python3
"""Run non-sensitive local parsing/OCR baseline checks for AuditPilot."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
import sys
import tempfile
import time
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile


ROOT = Path(__file__).resolve().parents[1]
BACKEND_DIR = ROOT / "backend"
if str(BACKEND_DIR) not in sys.path:
    sys.path.insert(0, str(BACKEND_DIR))

from app.services.document_parser import DocumentParser  # noqa: E402
from app.services.ocr_service import OCRService  # noqa: E402


class SyntheticOCR:
    def __init__(self) -> None:
        self.calls = 0

    def extract_text(self, image_path: Path) -> str:
        self.calls += 1
        return f"SYNTHETIC_OCR_PAGE_{self.calls}"


def _write_text(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def _write_image(path: Path) -> None:
    try:
        from PIL import Image, ImageDraw
    except Exception:
        path.write_bytes(b"synthetic-image")
        return

    image = Image.new("RGB", (360, 140), "white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 50), "AuditPilot 123", fill="black")
    image.save(path)


def _write_pdf(path: Path, pages: int, text: str | None) -> None:
    import fitz

    document = fitz.open()
    for index in range(pages):
        page = document.new_page()
        if text:
            page.insert_text((72, 72), f"{text} page {index + 1}")
    document.save(path)
    document.close()


def _write_docx(path: Path) -> None:
    try:
        from docx import Document
    except Exception:
        _write_minimal_docx(path)
        return

    document = Document()
    document.add_paragraph("AuditPilot DOCX synthetic procurement material")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Control"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Approval"
    table.cell(1, 1).text = "Complete"
    document.save(path)


def _write_xlsx(path: Path) -> None:
    try:
        from openpyxl import Workbook
    except Exception:
        _write_minimal_xlsx(path)
        return

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Procurement"
    sheet.append(["Project", "Amount", "Status"])
    sheet.append(["Synthetic", 100, "Approved"])
    workbook.save(path)


def _write_minimal_docx(path: Path) -> None:
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            (
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                "</Types>"
            ),
        )
        archive.writestr(
            "_rels/.rels",
            (
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
                'Target="word/document.xml"/>'
                "</Relationships>"
            ),
        )
        archive.writestr(
            "word/document.xml",
            (
                '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                "<w:body><w:p><w:r><w:t>AuditPilot DOCX synthetic material</w:t></w:r></w:p></w:body>"
                "</w:document>"
            ),
        )


def _write_minimal_xlsx(path: Path) -> None:
    with ZipFile(path, "w", ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            (
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/xl/workbook.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
                '<Override PartName="/xl/worksheets/sheet1.xml" '
                'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
                "</Types>"
            ),
        )
        archive.writestr(
            "xl/workbook.xml",
            (
                '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" '
                'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                '<sheets><sheet name="Procurement" sheetId="1" r:id="rId1"/></sheets></workbook>'
            ),
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            (
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" '
                'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
                'Target="worksheets/sheet1.xml"/></Relationships>'
            ),
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            (
                '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
                "<sheetData><row><c t=\"inlineStr\"><is><t>AuditPilot XLSX synthetic</t></is></c></row></sheetData>"
                "</worksheet>"
            ),
        )


def _measure(parser: DocumentParser, path: Path, name: str, pages: int | None = None) -> dict[str, Any]:
    start = time.perf_counter()
    parsed = parser.parse(path, path.name)
    elapsed = time.perf_counter() - start
    return {
        "file_type": name,
        "filename": path.name,
        "file_size_bytes": path.stat().st_size,
        "pages": pages,
        "parse_status": parsed.status,
        "parse_elapsed_seconds": round(elapsed, 4),
        "ocr_elapsed_seconds": None,
        "llm_elapsed_seconds": None,
        "total_elapsed_seconds": round(elapsed, 4),
        "error": parsed.error,
    }


def run(real_ocr: bool) -> list[dict[str, Any]]:
    with tempfile.TemporaryDirectory(prefix="auditpilot_baseline_") as temp_dir:
        temp = Path(temp_dir)
        image = temp / "single_image.png"
        _write_image(image)
        text_pdf = temp / "text.pdf"
        scan_pdf_1 = temp / "scan_1.pdf"
        scan_pdf_10 = temp / "scan_10.pdf"
        docx = temp / "sample.docx"
        xlsx = temp / "sample.xlsx"
        _write_pdf(text_pdf, 1, "Text PDF synthetic procurement control " * 3)
        _write_pdf(scan_pdf_1, 1, None)
        _write_pdf(scan_pdf_10, 10, None)
        _write_docx(docx)
        _write_xlsx(xlsx)

        ocr_service = OCRService() if real_ocr else SyntheticOCR()
        parser = DocumentParser(ocr_service=ocr_service)
        results = [
            _measure(parser, image, "single_image_ocr"),
            _measure(parser, text_pdf, "text_pdf", pages=1),
            _measure(parser, scan_pdf_1, "single_page_scanned_pdf", pages=1),
            _measure(parser, scan_pdf_10, "ten_page_scanned_pdf", pages=10),
            _measure(parser, docx, "docx"),
            _measure(parser, xlsx, "xlsx"),
        ]

        batch_start = time.perf_counter()
        for index in range(5):
            batch_image = temp / f"batch_{index}.png"
            _write_image(batch_image)
            parser.parse(batch_image, batch_image.name)
        batch_elapsed = time.perf_counter() - batch_start
        results.append(
            {
                "file_type": "five_image_batch",
                "filename": "batch_*.png",
                "file_size_bytes": None,
                "pages": None,
                "parse_status": "parsed",
                "parse_elapsed_seconds": round(batch_elapsed, 4),
                "ocr_elapsed_seconds": None,
                "llm_elapsed_seconds": None,
                "total_elapsed_seconds": round(batch_elapsed, 4),
                "error": None,
            }
        )

        mixed_start = time.perf_counter()
        for path in [image, text_pdf, scan_pdf_1, docx, xlsx]:
            parser.parse(path, path.name)
        mixed_elapsed = time.perf_counter() - mixed_start
        results.append(
            {
                "file_type": "mixed_upload_parse_loop",
                "filename": "image+pdf+docx+xlsx",
                "file_size_bytes": None,
                "pages": None,
                "parse_status": "parsed",
                "parse_elapsed_seconds": round(mixed_elapsed, 4),
                "ocr_elapsed_seconds": None,
                "llm_elapsed_seconds": None,
                "total_elapsed_seconds": round(mixed_elapsed, 4),
                "error": None,
            }
        )
        return results


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--real-ocr", action="store_true", help="use configured real OCR provider")
    parser.add_argument("--output", type=Path, help="optional JSON output path")
    args = parser.parse_args()

    results = run(real_ocr=args.real_ocr)
    text = json.dumps(results, ensure_ascii=False, indent=2)
    if args.output:
        args.output.write_text(text + "\n", encoding="utf-8")
    print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
